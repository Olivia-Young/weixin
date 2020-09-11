# 2020年8月29日
# 本次目标利用搜狗微信网站，进行指定公众号文章最新内容爬取。


import requests
from urllib.parse import quote
from pyquery import PyQuery as pq
import time
from docx import Document
from datetime import datetime
import random

header = [
    'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/84.0.4147.105 Safari/537.36',
    'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.25 Safari/537.36 Core/1.70.3775.400 QQBrowser',

]

#  该函数获取结果页面内容
def get_page(page,name):
    s = quote(name)
    headers = {
            'User-Agent':random.choice(header),
        }
    url = 'https://weixin.sogou.com/weixin?type=2&page={}&s_from=input&query={}&ie=utf8&_sug_=n&_sug_type_='.format(page,s)
    a = requests.get(url,headers=headers)

    return a.text

# 找出对应公众号的文章链接和标题，因为搜狗微信只显示前10页，所以我们也只抓取前10页，找出两天类的文章并输出
def get_article(name):
    for i in range(1,11):
        q = pq(get_page(i,name))
        lists = q('.news-list li').items()
        for list in lists:
            if list('.s-p a').text() == name:

                # 获取文章的时间戳
                t = list('.s-p').attr('t')
                # 这里是计算当前的时间戳
                shijian = int(time.time())

                # 这里是将两天内的文章搜索出来,并返回

                if (shijian - int(t)) < 172800:
                    title = list('.txt-box h3 a ').text()  # 获取文章标题
                    if title == '':
                        title = '无最新报价'
                    url = 'https://weixin.sogou.com' + list('.txt-box h3 a ').attr('href')  # 获取文章链接
                    if url == '':
                        url = ''

                    print(title,url)
                    return title,url


    # 上面函数已经拿到文章的名字和文章链接，现在来进行文章内容分析，（这块代码放弃了，因为公众号正文是反爬的加入cookie和selenium都会被劫住，所以暂时不搞这一块）
    # def get_content(self,title,url):
    #     brower = webdriver.Chrome()
    #
    #     time.sleep(2)
    #     print(url)
    #     brower.get(url)
    #     html = brower.page_source
    #     h = pq(html)
    #     print(h)
    #
    #     # lists = h('.rich_media_content p').items()
    #     # for list in lists:
    #     #
    #     #     print(list.text())


# 这里我们建立一个函数完成所有的文章标题链接写入到word文档中
lists = ['西安腾乐电子','西安新志电子']
def execute():
    # 创建一个空文档
    document = Document()
    #给文档添加一个标题,后面加上时间
    header = '公众号最新报价单({})'.format(datetime.now().strftime('%a,%b-%d %H:%M'))
    document.add_heading(header,level=0)
    for list in lists:
        print(list)
        time.sleep(3)
        title,url = get_article(list)
        if title == '':
            continue
        document.add_paragraph(title,style='List Number')
        document.add_paragraph('链接：' + url + '\n')
        document.add_paragraph('来自：' + list)

    document.save('最新报价.docx')



if __name__ == '__main__':
    execute()


